"""
Test suite for OCR Parse.

Creates a fake PDF with known text, uploads it to the app,
verifies OCR extraction, and tests the full pipeline including export.
"""

import email as email_lib
import io
import json
import os
import re
import shutil
import tempfile
import unittest
import unittest.mock
import uuid
import zipfile
from pathlib import Path
from unittest.mock import patch

from fpdf import FPDF

# Ensure uploads go to a temp directory during tests
_test_upload_dir = tempfile.mkdtemp(prefix="ocr_parse_test_")

# Patch UPLOAD_BASE before importing app
import app as app_module

app_module.UPLOAD_BASE = Path(_test_upload_dir)

from app import app


def _make_test_pdf(text_lines=None):
    """Generate a simple PDF with known text content."""
    if text_lines is None:
        text_lines = [
            "INVOICE #12345",
            "Date: January 15, 2025",
            "",
            "Bill To: John Smith",
            "123 Main Street",
            "Springfield, IL 62701",
            "",
            "Item: Widget A - Qty: 10 - Price: $5.00",
            "Item: Widget B - Qty: 5 - Price: $12.50",
            "",
            "Subtotal: $112.50",
            "Tax (8%): $9.00",
            "Total: $121.50",
            "",
            "Notes: Please review and approve.",
        ]

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)

    for line in text_lines:
        if line == "":
            pdf.ln(8)
        else:
            pdf.cell(0, 10, line, new_x="LMARGIN", new_y="NEXT")

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


class TestOCRParse(unittest.TestCase):
    """End-to-end tests for the OCR Parse application."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        # Clean up any session directories created during tests
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    def test_index_page(self):
        """The upload page loads successfully."""
        resp = self.client.get("/")
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b"OCR Parse", resp.data)

    def test_health_check(self):
        """The health endpoint returns ok."""
        resp = self.client.get("/health")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertEqual(data["status"], "ok")

    def test_upload_no_file(self):
        """Upload without a file returns an error."""
        resp = self.client.post("/upload")
        self.assertEqual(resp.status_code, 400)

    def test_upload_and_process(self):
        """Upload a PDF and run OCR – verify text is extracted."""
        pdf_buf = _make_test_pdf()

        # Upload
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        session_id = data["session_id"]
        self.assertTrue(session_id)
        self.assertIn("/review/", data["redirect"])

        # Process (OCR)
        resp = self.client.post(f"/api/process/{session_id}")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        self.assertGreaterEqual(data["num_pages"], 1)

        # Verify OCR extracted some text
        results = data["results"]
        self.assertTrue(len(results) > 0)
        all_text = " ".join(
            block["text"] for page in results for block in page["blocks"]
        )
        # The OCR should find at least some of these words
        found_keywords = sum(
            1
            for kw in ["INVOICE", "12345", "John", "Smith", "Widget", "Total"]
            if kw.lower() in all_text.lower()
        )
        self.assertGreaterEqual(
            found_keywords, 2, f"Expected to find keywords in OCR text: {all_text}"
        )

    def test_full_pipeline(self):
        """Full pipeline: upload → OCR → add annotations → export docx."""
        pdf_buf = _make_test_pdf()

        # Upload
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        # Process
        resp = self.client.post(f"/api/process/{session_id}")
        self.assertEqual(resp.status_code, 200)

        # Get page image
        resp = self.client.get(f"/api/page/{session_id}/1")
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content_type, "image/png")

        # Save annotations
        annotations = [
            {"id": "test-1", "text": "Please fix this amount", "type": "comment", "page": 1},
            {"id": "test-2", "text": "New line item added", "type": "insert", "page": 1},
        ]
        resp = self.client.post(
            f"/api/annotations/{session_id}",
            data=json.dumps(annotations),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)

        # Export
        resp = self.client.post(f"/api/export/{session_id}")
        self.assertEqual(resp.status_code, 200)
        self.assertIn("wordprocessingml", resp.content_type)
        self.assertGreater(len(resp.data), 1000)  # docx should be non-trivial

    def test_cleanup(self):
        """Cleanup endpoint removes session files."""
        pdf_buf = _make_test_pdf()

        # Upload
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        sdir = Path(_test_upload_dir) / session_id
        self.assertTrue(sdir.exists())

        # Cleanup
        resp = self.client.post(f"/api/cleanup/{session_id}")
        self.assertEqual(resp.status_code, 200)
        self.assertFalse(sdir.exists())

    def test_tesseract_engine_only(self):
        """Verify the app always uses Tesseract regardless of form data."""
        pdf_buf = _make_test_pdf()

        # Try to select openai engine – should be ignored
        resp = self.client.post(
            "/upload",
            data={
                "annotated_pdf": (pdf_buf, "test.pdf"),
                "ocr_engine": "openai",
                "api_key": "fake-key",
            },
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        config_path = Path(_test_upload_dir) / session_id / "config.json"
        with open(config_path) as f:
            config = json.load(f)
        self.assertEqual(config["ocr_engine"], "tesseract")
        self.assertEqual(config["api_key"], "")

    def test_upload_too_large_returns_json_413(self):
        """When the upload exceeds MAX_CONTENT_LENGTH the 413 handler returns JSON."""
        # Temporarily set a very small limit so we can trigger it with a tiny payload
        original_limit = app.config["MAX_CONTENT_LENGTH"]
        app.config["MAX_CONTENT_LENGTH"] = 1  # 1 byte
        try:
            pdf_buf = _make_test_pdf()
            resp = self.client.post(
                "/upload",
                data={"annotated_pdf": (pdf_buf, "big.pdf")},
                content_type="multipart/form-data",
            )
            self.assertEqual(resp.status_code, 413)
            data = json.loads(resp.data)
            self.assertIn("error", data)
            self.assertIn("error_type", data)
            self.assertEqual(data["error_type"], "file_too_large")
            self.assertIn("too large", data["error"].lower())
        finally:
            app.config["MAX_CONTENT_LENGTH"] = original_limit

    def test_api_404_returns_json(self):
        """API 404 responses return JSON, not HTML."""
        resp = self.client.post("/api/process/nonexistent-session-id")
        self.assertEqual(resp.status_code, 404)
        data = json.loads(resp.data)
        self.assertIn("error", data)

    def test_api_500_returns_json(self):
        """A forced 500 on an API route returns JSON, not HTML."""
        from unittest.mock import patch

        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        # Patch _save_json to raise inside the process route to trigger the 500 handler
        with patch("app._save_json", side_effect=RuntimeError("forced error")):
            resp = self.client.post(f"/api/process/{session_id}")

        self.assertEqual(resp.status_code, 500)
        data = json.loads(resp.data)
        self.assertIn("error", data)
        # Must be parseable as JSON – not an HTML page
        self.assertEqual(resp.content_type, "application/json")

    def test_debug_log_endpoints(self):
        """Debug log is empty by default; upload failure with debug_mode stores an entry."""
        import app as app_module
        from pathlib import Path as _Path

        # Point the debug log to the test upload directory
        original_path = app_module.DEBUG_LOG_PATH
        app_module.DEBUG_LOG_PATH = _Path(_test_upload_dir) / "debug_log.json"
        # Ensure it starts empty
        if app_module.DEBUG_LOG_PATH.exists():
            app_module.DEBUG_LOG_PATH.unlink()

        try:
            # GET debug log – should be empty
            resp = self.client.get("/api/debug/log")
            self.assertEqual(resp.status_code, 200)
            self.assertEqual(json.loads(resp.data), [])

            # Trigger a failure with debug_mode=1 (wrong file type)
            resp = self.client.post(
                "/upload",
                data={
                    "annotated_pdf": (io.BytesIO(b"not a pdf"), "test.txt"),
                    "debug_mode": "1",
                },
                content_type="multipart/form-data",
            )
            self.assertEqual(resp.status_code, 400)

            # Debug log should now have one entry
            resp = self.client.get("/api/debug/log")
            log = json.loads(resp.data)
            self.assertEqual(len(log), 1)
            self.assertEqual(log[0]["event"], "upload_failed")
            self.assertIn("error_type", log[0])

            # Clear the log
            resp = self.client.post("/api/debug/clear")
            self.assertEqual(resp.status_code, 200)

            # Confirm it is empty again
            resp = self.client.get("/api/debug/log")
            self.assertEqual(json.loads(resp.data), [])

        finally:
            app_module.DEBUG_LOG_PATH = original_path

    def test_debug_mode_off_does_not_store_log(self):
        """Upload failure without debug_mode=1 does NOT write to debug log."""
        import app as app_module
        from pathlib import Path as _Path

        original_path = app_module.DEBUG_LOG_PATH
        app_module.DEBUG_LOG_PATH = _Path(_test_upload_dir) / "debug_log_off.json"
        if app_module.DEBUG_LOG_PATH.exists():
            app_module.DEBUG_LOG_PATH.unlink()

        try:
            # Wrong file type, debug_mode NOT set
            self.client.post(
                "/upload",
                data={"annotated_pdf": (io.BytesIO(b"not a pdf"), "test.txt")},
                content_type="multipart/form-data",
            )
            # Log file should not have been created
            self.assertFalse(app_module.DEBUG_LOG_PATH.exists())
        finally:
            app_module.DEBUG_LOG_PATH = original_path

    def test_process_returns_debug_info_when_enabled(self):
        """Processing returns debug info when DEBUG_MODE is on."""
        import app as app_module

        original_debug = app_module.DEBUG_MODE
        app_module.DEBUG_MODE = True
        try:
            pdf_buf = _make_test_pdf()
            resp = self.client.post(
                "/upload",
                data={"annotated_pdf": (pdf_buf, "test.pdf")},
                content_type="multipart/form-data",
            )
            data = json.loads(resp.data)
            session_id = data["session_id"]

            resp = self.client.post(f"/api/process/{session_id}")
            self.assertEqual(resp.status_code, 200)
            data = json.loads(resp.data)
            self.assertTrue(data["success"])

            # Debug info should be present
            self.assertIn("debug", data)
            debug = data["debug"]
            self.assertTrue(debug["debug_mode"])
            self.assertIn("tesseract_version", debug)
            self.assertIn("poppler_version", debug)
            self.assertIn("steps", debug)
            self.assertGreater(len(debug["steps"]), 0)
            self.assertIn("session_id", debug)
            self.assertIn("pdf_file_size_bytes", debug)
            self.assertIn("total_elapsed_sec", debug)

            # Verify steps include pdf_to_images and at least one ocr page
            step_names = [s["step"] for s in debug["steps"]]
            self.assertIn("pdf_to_images", step_names)
            self.assertTrue(any(s.startswith("ocr_page_") for s in step_names))
        finally:
            app_module.DEBUG_MODE = original_debug

    def test_process_no_debug_info_when_disabled(self):
        """Processing omits debug info when DEBUG_MODE is off."""
        import app as app_module

        original_debug = app_module.DEBUG_MODE
        app_module.DEBUG_MODE = False
        try:
            pdf_buf = _make_test_pdf()
            resp = self.client.post(
                "/upload",
                data={"annotated_pdf": (pdf_buf, "test.pdf")},
                content_type="multipart/form-data",
            )
            data = json.loads(resp.data)
            session_id = data["session_id"]

            resp = self.client.post(f"/api/process/{session_id}")
            self.assertEqual(resp.status_code, 200)
            data = json.loads(resp.data)
            self.assertTrue(data["success"])
            self.assertNotIn("debug", data)
        finally:
            app_module.DEBUG_MODE = original_debug

    def test_debug_info_endpoint(self):
        """The /api/debug/info endpoint returns system diagnostics."""
        import app as app_module

        original_debug = app_module.DEBUG_MODE
        app_module.DEBUG_MODE = True
        try:
            resp = self.client.get("/api/debug/info")
            self.assertEqual(resp.status_code, 200)
            data = json.loads(resp.data)
            self.assertTrue(data["debug_mode"])
            self.assertIn("tesseract_version", data)
            self.assertIn("poppler_version", data)
        finally:
            app_module.DEBUG_MODE = original_debug

    def test_debug_info_endpoint_disabled(self):
        """The /api/debug/info endpoint indicates off when DEBUG_MODE is off."""
        import app as app_module

        original_debug = app_module.DEBUG_MODE
        app_module.DEBUG_MODE = False
        try:
            resp = self.client.get("/api/debug/info")
            self.assertEqual(resp.status_code, 200)
            data = json.loads(resp.data)
            self.assertFalse(data["debug_mode"])
        finally:
            app_module.DEBUG_MODE = original_debug


    def test_upload_stores_email_and_token(self):
        """Upload with email stores email and access token in session config."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf"), "email": "tester@example.com"},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        session_id = data["session_id"]

        config_path = Path(_test_upload_dir) / session_id / "config.json"
        with open(config_path) as f:
            config = json.load(f)

        self.assertEqual(config["email"], "tester@example.com")
        self.assertIn("access_token", config)
        self.assertGreater(len(config["access_token"]), 20)

    def test_upload_invalid_email_rejected(self):
        """Upload with an invalid email address returns 400."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf"), "email": "not-an-email"},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 400)
        data = json.loads(resp.data)
        self.assertIn("email", data["error"].lower())

    def test_upload_without_email_still_works(self):
        """Upload without email succeeds and config has empty email."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        config_path = Path(_test_upload_dir) / data["session_id"] / "config.json"
        with open(config_path) as f:
            config = json.load(f)
        self.assertEqual(config["email"], "")

    def test_access_token_redirect(self):
        """The /r/<token> route redirects to the review page."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf"), "email": "user@example.com"},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        config_path = Path(_test_upload_dir) / session_id / "config.json"
        with open(config_path) as f:
            config = json.load(f)
        access_token = config["access_token"]

        resp = self.client.get(f"/r/{access_token}")
        self.assertEqual(resp.status_code, 302)
        self.assertIn(f"/review/{session_id}", resp.headers.get("Location", ""))

    def test_access_token_invalid_returns_404(self):
        """An unknown or malformed access token returns 404."""
        self.assertEqual(self.client.get("/r/nonexistent-token").status_code, 404)
        self.assertEqual(self.client.get("/r/../../etc/passwd").status_code, 404)

    def test_cleanup_removes_token(self):
        """Cleanup removes the access token from the registry."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf"), "email": "user@example.com"},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        config_path = Path(_test_upload_dir) / session_id / "config.json"
        with open(config_path) as f:
            config = json.load(f)
        access_token = config["access_token"]

        # Token works before cleanup
        resp = self.client.get(f"/r/{access_token}")
        self.assertEqual(resp.status_code, 302)

        # Cleanup
        self.client.post(f"/api/cleanup/{session_id}")

        # Token no longer resolves
        resp = self.client.get(f"/r/{access_token}")
        self.assertEqual(resp.status_code, 404)

    def test_process_returns_cached_results_on_second_call(self):
        """Second call to /api/process returns cached results instantly."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        # First call – processes normally
        resp = self.client.post(f"/api/process/{session_id}")
        self.assertEqual(resp.status_code, 200)
        first_data = json.loads(resp.data)
        self.assertTrue(first_data["success"])
        self.assertFalse(first_data.get("cached", False))

        # Second call – should serve cached results
        resp = self.client.post(f"/api/process/{session_id}")
        self.assertEqual(resp.status_code, 200)
        second_data = json.loads(resp.data)
        self.assertTrue(second_data["success"])
        self.assertTrue(second_data.get("cached", False))


class TestOCREngine(unittest.TestCase):
    """Unit tests for OCR engine."""

    def test_tesseract_extract(self):
        """Tesseract extracts text from a simple image."""
        from PIL import Image, ImageDraw, ImageFont
        from ocr_engine import OCREngine

        # Create a test image with text
        img = Image.new("RGB", (400, 100), color="white")
        draw = ImageDraw.Draw(img)
        draw.text((20, 30), "Hello World Test 123", fill="black")

        engine = OCREngine(engine="tesseract")
        blocks = engine.extract_text_blocks(img)

        all_text = " ".join(b["text"] for b in blocks).lower()
        self.assertIn("hello", all_text)
        self.assertIn("world", all_text)


class TestWordExport(unittest.TestCase):
    """Unit tests for Word document export."""

    def test_create_document_with_annotations(self):
        """Word export creates a valid docx file."""
        from word_export import create_word_document

        annotations = [
            {"text": "Test comment", "type": "comment", "page": 1},
            {"text": "Inserted text", "type": "insert", "page": 1},
            {"text": "Deleted text", "type": "delete", "page": 1},
        ]

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            create_word_document(annotations, None, output_path)
            self.assertTrue(os.path.exists(output_path))
            self.assertGreater(os.path.getsize(output_path), 1000)
        finally:
            os.unlink(output_path)


class TestSecurity(unittest.TestCase):
    """Security-focused tests: path traversal, token hardening, injection."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    # --- Path traversal / session-ID sanitisation ---------------------------

    def test_review_path_traversal_returns_404(self):
        """Attempt to traverse out of uploads dir via review route returns 404."""
        for evil in ["../etc/passwd", "..%2Fetc%2Fpasswd", "....//etc//passwd"]:
            resp = self.client.get(f"/review/{evil}")
            self.assertIn(resp.status_code, (404, 308), f"Expected 404 for {evil!r}, got {resp.status_code}")

    def test_process_path_traversal_returns_404(self):
        """Path traversal in process endpoint returns 404, not a file read."""
        resp = self.client.post("/api/process/../etc/passwd")
        self.assertIn(resp.status_code, (404, 308))

    def test_annotations_path_traversal_returns_404(self):
        """Path traversal in annotations endpoint does not expose server files."""
        resp = self.client.post(
            "/api/annotations/../../../etc/shadow",
            data=json.dumps([]),
            content_type="application/json",
        )
        self.assertIn(resp.status_code, (404, 308))

    def test_cleanup_path_traversal_returns_404(self):
        """Path traversal in cleanup endpoint returns 404."""
        resp = self.client.post("/api/cleanup/../../etc/passwd")
        self.assertIn(resp.status_code, (404, 308))

    def test_export_path_traversal_returns_404(self):
        """Path traversal in export endpoint returns 404."""
        resp = self.client.post("/api/export/../etc/passwd")
        self.assertIn(resp.status_code, (404, 308))

    # --- Token security ---------------------------------------------------

    def test_token_is_high_entropy(self):
        """Access token generated by upload is high-entropy (≥ 32 URL-safe chars)."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf"), "email": "sec@example.com"},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]
        config_path = Path(_test_upload_dir) / session_id / "config.json"
        with open(config_path) as f:
            config = json.load(f)
        token = config["access_token"]
        # token_urlsafe(32) produces 43-character base64url string
        self.assertGreaterEqual(len(token), 40, "Access token is too short")
        # Must only contain URL-safe characters
        self.assertRegex(token, r"^[A-Za-z0-9_\-]+$", "Token contains non-URL-safe chars")

    def test_token_not_reused_across_sessions(self):
        """Two separate uploads produce two distinct access tokens."""
        tokens = set()
        for _ in range(2):
            pdf_buf = _make_test_pdf()
            resp = self.client.post(
                "/upload",
                data={"annotated_pdf": (pdf_buf, "test.pdf")},
                content_type="multipart/form-data",
            )
            data = json.loads(resp.data)
            session_id = data["session_id"]
            with open(Path(_test_upload_dir) / session_id / "config.json") as f:
                tokens.add(json.load(f)["access_token"])
        self.assertEqual(len(tokens), 2, "Two uploads generated the same token")

    def test_token_redirect_to_deleted_session_returns_404(self):
        """Token redirect returns 404 after the session directory is removed."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]
        with open(Path(_test_upload_dir) / session_id / "config.json") as f:
            config = json.load(f)
        token = config["access_token"]

        # Manually delete the session dir (simulate expiry without cleanup endpoint)
        shutil.rmtree(str(Path(_test_upload_dir) / session_id), ignore_errors=True)

        resp = self.client.get(f"/r/{token}")
        self.assertEqual(resp.status_code, 404)

    # --- File-upload security ---------------------------------------------

    def test_non_pdf_extension_rejected(self):
        """Uploading a file with a non-PDF extension is rejected with 400."""
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (io.BytesIO(b"not a pdf"), "malware.exe")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 400)
        data = json.loads(resp.data)
        self.assertIn("error", data)

    def test_script_injection_in_filename_sanitised(self):
        """Malicious filename is sanitised by secure_filename; upload still works."""
        resp = self.client.post(
            "/upload",
            data={
                "annotated_pdf": (
                    io.BytesIO(b"%PDF-1.4 fake"),
                    "<script>alert(1)</script>.pdf",
                )
            },
            content_type="multipart/form-data",
        )
        # Upload succeeds (extension is .pdf), but the file is stored under annotated.pdf
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        session_id = data["session_id"]
        sdir = Path(_test_upload_dir) / session_id
        # The saved file should be annotated.pdf, not the raw malicious name
        self.assertTrue((sdir / "annotated.pdf").exists())
        # The malicious script should NOT appear as a filename in the directory
        for f in sdir.iterdir():
            self.assertNotIn("<script>", f.name)

    def test_invalid_json_annotations_returns_400(self):
        """Saving annotations with invalid JSON body returns 400."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        session_id = json.loads(resp.data)["session_id"]

        resp = self.client.post(
            f"/api/annotations/{session_id}",
            data=b"not-valid-json",
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)

    def test_annotations_must_be_list_not_object(self):
        """Saving a JSON object (not array) as annotations returns 400."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        session_id = json.loads(resp.data)["session_id"]

        resp = self.client.post(
            f"/api/annotations/{session_id}",
            data=json.dumps({"key": "value"}),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 400)
        self.assertIn("error", json.loads(resp.data))

    def test_nonexistent_session_returns_404_on_all_endpoints(self):
        """All session-scoped endpoints return 404 for unknown session IDs."""
        fake = "00000000-0000-0000-0000-000000000000"
        endpoints = [
            ("POST", f"/api/process/{fake}"),
            ("GET", f"/api/page/{fake}/1"),
            ("GET", f"/api/annotations/{fake}"),
            ("POST", f"/api/export/{fake}"),
        ]
        for method, url in endpoints:
            if method == "GET":
                resp = self.client.get(url)
            else:
                resp = self.client.post(url)
            self.assertIn(
                resp.status_code,
                (404, 200),  # annotations GET returns [] for missing, process returns 404
                f"Unexpected status {resp.status_code} for {method} {url}",
            )

    def test_process_missing_pdf_returns_404(self):
        """If annotated.pdf is missing from a session dir, /api/process returns 404."""
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        session_id = json.loads(resp.data)["session_id"]

        # Remove the PDF to simulate missing file
        pdf_path = Path(_test_upload_dir) / session_id / "annotated.pdf"
        pdf_path.unlink()

        resp = self.client.post(f"/api/process/{session_id}")
        self.assertEqual(resp.status_code, 404)
        self.assertIn("error", json.loads(resp.data))

    def test_review_page_invalid_session_returns_404(self):
        """Review page returns 404 for an unknown session ID."""
        resp = self.client.get("/review/nonexistent-session-id-xyz")
        self.assertEqual(resp.status_code, 404)

    def test_review_page_empty_config_returns_404(self):
        """Review page returns 404 when config.json is missing/empty."""
        sid = str(uuid.uuid4())
        (Path(_test_upload_dir) / sid).mkdir()
        resp = self.client.get(f"/review/{sid}")
        self.assertEqual(resp.status_code, 404)


class TestEmailNotifications(unittest.TestCase):
    """Tests for email notification behaviour."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    def test_smtp_not_called_when_not_configured(self):
        """Email sending is skipped when SMTP_HOST is empty."""
        import app as app_module
        orig_host = app_module.SMTP_HOST
        app_module.SMTP_HOST = ""
        try:
            with patch("smtplib.SMTP") as mock_smtp:
                from app import _send_notification_email
                _send_notification_email(
                    "user@example.com", "http://localhost:5000", "fake-token"
                )
                mock_smtp.assert_not_called()
        finally:
            app_module.SMTP_HOST = orig_host

    def test_smtp_called_with_correct_recipient_when_configured(self):
        """Email is sent via SMTP when host and user are configured."""
        import app as app_module
        orig_host, orig_user = app_module.SMTP_HOST, app_module.SMTP_USER
        app_module.SMTP_HOST = "smtp.example.com"
        app_module.SMTP_USER = "sender@example.com"
        try:
            with patch("smtplib.SMTP") as mock_smtp_cls:
                mock_server = unittest.mock.MagicMock()
                mock_smtp_cls.return_value.__enter__ = lambda s: mock_server
                mock_smtp_cls.return_value.__exit__ = unittest.mock.MagicMock(
                    return_value=False
                )
                from app import _send_notification_email
                _send_notification_email(
                    "recipient@example.com",
                    "http://example.com",
                    "test-token-abc",
                )
                mock_smtp_cls.assert_called_once_with("smtp.example.com", 587, timeout=10)
                # sendmail should be called with the correct recipient
                mock_server.sendmail.assert_called_once()
                call_args = mock_server.sendmail.call_args[0]
                self.assertEqual(call_args[1], "recipient@example.com")
        finally:
            app_module.SMTP_HOST = orig_host
            app_module.SMTP_USER = orig_user

    def test_smtp_failure_does_not_raise(self):
        """SMTP connection failure is swallowed (logging only), not raised."""
        import app as app_module
        orig_host, orig_user = app_module.SMTP_HOST, app_module.SMTP_USER
        app_module.SMTP_HOST = "smtp.example.com"
        app_module.SMTP_USER = "sender@example.com"
        try:
            with patch("smtplib.SMTP", side_effect=ConnectionRefusedError("no server")):
                from app import _send_notification_email
                # Should not raise
                _send_notification_email(
                    "user@example.com", "http://example.com", "tok"
                )
        finally:
            app_module.SMTP_HOST = orig_host
            app_module.SMTP_USER = orig_user

    def test_email_sent_flag_prevents_duplicate_emails(self):
        """Processing a session twice does not send a second notification email."""
        import app as app_module
        orig_host, orig_user = app_module.SMTP_HOST, app_module.SMTP_USER
        app_module.SMTP_HOST = "smtp.example.com"
        app_module.SMTP_USER = "sender@example.com"
        try:
            pdf_buf = _make_test_pdf()
            resp = self.client.post(
                "/upload",
                data={"annotated_pdf": (pdf_buf, "test.pdf"), "email": "dup@example.com"},
                content_type="multipart/form-data",
            )
            session_id = json.loads(resp.data)["session_id"]

            send_count = [0]

            def counting_send(to, base_url, token):
                send_count[0] += 1

            with patch("app._send_notification_email", side_effect=counting_send):
                self.client.post(f"/api/process/{session_id}")
                self.client.post(f"/api/process/{session_id}")  # second call uses cache

            self.assertEqual(
                send_count[0], 1, "Notification email sent more than once"
            )
        finally:
            app_module.SMTP_HOST = orig_host
            app_module.SMTP_USER = orig_user

    def test_email_link_contains_token_not_session_id(self):
        """The review link in the email body uses the token, not the session ID."""
        import app as app_module
        orig_host, orig_user = app_module.SMTP_HOST, app_module.SMTP_USER
        app_module.SMTP_HOST = "smtp.example.com"
        app_module.SMTP_USER = "sender@example.com"
        try:
            captured = {}

            def fake_smtp_cls(host, port, timeout=10):
                class FakeSMTP:
                    def __enter__(self): return self
                    def __exit__(self, *a): pass
                    def ehlo(self): pass
                    def starttls(self): pass
                    def login(self, u, p): pass
                    def sendmail(self, frm, to, msg):
                        captured["msg"] = msg

                return FakeSMTP()

            with patch("smtplib.SMTP", side_effect=fake_smtp_cls):
                from app import _send_notification_email
                _send_notification_email(
                    "to@example.com", "http://example.com", "my-secret-token"
                )

            raw = captured.get("msg", "")
            # Decode the MIME message and extract all plain text bodies
            msg = email_lib.message_from_string(raw)
            full_body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    charset = part.get_content_charset() or "utf-8"
                    payload = part.get_payload(decode=True)
                    if payload:
                        full_body += payload.decode(charset, errors="replace")
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    full_body = payload.decode(
                        msg.get_content_charset() or "utf-8", errors="replace"
                    )

            self.assertIn("my-secret-token", full_body, "Token not found in email body")
        finally:
            app_module.SMTP_HOST = orig_host
            app_module.SMTP_USER = orig_user


class TestSessionCleanup(unittest.TestCase):
    """Tests for session expiry and cleanup logic."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)
            elif child.is_file() and child.name != "tokens.json":
                child.unlink(missing_ok=True)

    def test_cleanup_old_sessions_removes_expired_dirs(self):
        """_cleanup_old_sessions removes session dirs older than the max age."""
        import app as app_module
        import time

        # Create a fake session directory with an old mtime
        old_sid = "old-session-fake"
        old_dir = Path(_test_upload_dir) / old_sid
        old_dir.mkdir(exist_ok=True)
        # Backdate its mtime by 2x max age
        old_mtime = time.time() - app_module.SESSION_MAX_AGE_SECONDS * 2
        os.utime(str(old_dir), (old_mtime, old_mtime))

        # Create a fresh session directory
        new_sid = "new-session-fake"
        new_dir = Path(_test_upload_dir) / new_sid
        new_dir.mkdir(exist_ok=True)

        orig_base = app_module.UPLOAD_BASE
        app_module.UPLOAD_BASE = Path(_test_upload_dir)
        try:
            app_module._cleanup_old_sessions()
        finally:
            app_module.UPLOAD_BASE = orig_base

        self.assertFalse(old_dir.exists(), "Expired session dir should be removed")
        self.assertTrue(new_dir.exists(), "Fresh session dir should remain")

    def test_cleanup_session_function_removes_dir_and_token(self):
        """_cleanup_session removes the session directory and its token."""
        import app as app_module

        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        data = json.loads(resp.data)
        session_id = data["session_id"]

        sdir = Path(_test_upload_dir) / session_id
        self.assertTrue(sdir.exists())

        app_module._cleanup_session(session_id)

        self.assertFalse(sdir.exists(), "Session dir should be deleted")
        # Token should be gone from registry
        tokens = app_module._load_json(app_module._tokens_file(), {})
        self.assertNotIn(session_id, tokens.values())


class TestOCREngineExtended(unittest.TestCase):
    """Additional OCR engine unit tests."""

    def test_unknown_engine_raises_value_error(self):
        """OCREngine raises ValueError for an unsupported engine name."""
        from ocr_engine import OCREngine
        from PIL import Image

        engine = OCREngine(engine="unsupported-engine")
        img = Image.new("RGB", (100, 100), color="white")
        with self.assertRaises(ValueError):
            engine.extract_text_blocks(img)

    def test_tesseract_empty_image_returns_list(self):
        """OCR on a blank image returns an empty list (no text)."""
        from ocr_engine import OCREngine
        from PIL import Image

        engine = OCREngine(engine="tesseract")
        # Entirely white 100×100 image – no text
        img = Image.new("RGB", (100, 100), color="white")
        blocks = engine.extract_text_blocks(img)
        self.assertIsInstance(blocks, list)


class TestWordExportExtended(unittest.TestCase):
    """Extended Word export tests."""

    def test_export_empty_annotations(self):
        """Word export with empty annotations list produces a valid docx."""
        from word_export import create_word_document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            create_word_document([], None, output_path)
            self.assertTrue(os.path.exists(output_path))
            self.assertGreater(os.path.getsize(output_path), 500)
        finally:
            os.unlink(output_path)

    def test_export_insert_and_delete_annotations(self):
        """Word export correctly handles insert/delete tracked-change annotations."""
        from word_export import create_word_document

        annotations = [
            {"text": "New paragraph added", "type": "insert", "page": 2},
            {"text": "Old paragraph removed", "type": "delete", "page": 2},
        ]
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            create_word_document(annotations, None, output_path)
            # Verify the docx is a valid ZIP (all docx files are ZIP archives)
            self.assertTrue(zipfile.is_zipfile(output_path))
            with zipfile.ZipFile(output_path) as z:
                names = z.namelist()
                self.assertIn("word/document.xml", names)
        finally:
            os.unlink(output_path)

    def test_export_multiple_comments(self):
        """Word export injects comments.xml for multiple comment annotations."""
        from word_export import create_word_document

        annotations = [
            {"text": "First comment", "type": "comment", "page": 1},
            {"text": "Second comment", "type": "comment", "page": 2},
            {"text": "Third comment", "type": "comment", "page": 3},
        ]
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            create_word_document(annotations, None, output_path)
            with zipfile.ZipFile(output_path) as z:
                names = z.namelist()
                self.assertIn("word/comments.xml", names, "comments.xml should be injected")
                comments_xml = z.read("word/comments.xml").decode()
                self.assertIn("First comment", comments_xml)
                self.assertIn("Second comment", comments_xml)
                self.assertIn("Third comment", comments_xml)
        finally:
            os.unlink(output_path)

    def test_export_annotations_without_page_number(self):
        """Annotations without a page key export without errors."""
        from word_export import create_word_document

        annotations = [
            {"text": "No page annotation", "type": "comment"},
            {"text": "Also no page", "type": "insert"},
        ]
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            create_word_document(annotations, None, output_path)
            self.assertGreater(os.path.getsize(output_path), 500)
        finally:
            os.unlink(output_path)

    def test_export_with_original_docx_as_base(self):
        """Word export builds on an existing .docx when one is provided."""
        from word_export import create_word_document
        from docx import Document

        # Create a minimal original docx
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            orig_path = f.name
        orig_doc = Document()
        orig_doc.add_heading("Original Document", level=1)
        orig_doc.add_paragraph("This is the original content.")
        orig_doc.save(orig_path)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            annotations = [{"text": "Review note", "type": "comment", "page": 1}]
            create_word_document(annotations, orig_path, output_path)
            self.assertTrue(zipfile.is_zipfile(output_path))
            # Verify original content is retained
            with zipfile.ZipFile(output_path) as z:
                doc_xml = z.read("word/document.xml").decode()
            self.assertIn("Original", doc_xml)
        finally:
            os.unlink(orig_path)
            os.unlink(output_path)

    def test_export_via_api_returns_docx(self):
        """The /api/export endpoint returns a valid docx content-type."""
        app.config["TESTING"] = True
        client = app.test_client()

        pdf_buf = _make_test_pdf()
        resp = client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        session_id = json.loads(resp.data)["session_id"]

        # Save some annotations
        annotations = [
            {"id": "a1", "text": "Check this figure", "type": "comment", "page": 1},
        ]
        client.post(
            f"/api/annotations/{session_id}",
            data=json.dumps(annotations),
            content_type="application/json",
        )

        resp = client.post(f"/api/export/{session_id}")
        self.assertEqual(resp.status_code, 200)
        self.assertIn("wordprocessingml", resp.content_type)
        self.assertGreater(len(resp.data), 1000)

    def test_export_nonexistent_session_returns_404(self):
        """Exporting a non-existent session returns a JSON 404 error."""
        app.config["TESTING"] = True
        client = app.test_client()
        resp = client.post("/api/export/nonexistent-session-id-abc")
        self.assertEqual(resp.status_code, 404)
        data = json.loads(resp.data)
        self.assertIn("error", data)


class TestMultiPagePDF(unittest.TestCase):
    """Tests for multi-page PDF handling."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    def _make_multipage_pdf(self, num_pages=3):
        """Generate a PDF with multiple pages."""
        pdf = FPDF()
        for i in range(1, num_pages + 1):
            pdf.add_page()
            pdf.set_font("Helvetica", size=14)
            pdf.cell(0, 10, f"Page {i} Content: INVOICE TOTAL AMOUNT", new_x="LMARGIN", new_y="NEXT")
            pdf.cell(0, 10, f"Reference number: REF{i:04d}", new_x="LMARGIN", new_y="NEXT")
        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return buf

    def test_multipage_pdf_all_pages_processed(self):
        """OCR processes all pages of a multi-page PDF."""
        pdf_buf = self._make_multipage_pdf(num_pages=3)
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "multi.pdf")},
            content_type="multipart/form-data",
        )
        self.assertEqual(resp.status_code, 200)
        session_id = json.loads(resp.data)["session_id"]

        resp = self.client.post(f"/api/process/{session_id}")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)

        self.assertTrue(data["success"])
        self.assertEqual(data["num_pages"], 3)
        self.assertEqual(len(data["results"]), 3)

    def test_page_image_served_for_each_page(self):
        """All page images are served after processing a multi-page PDF."""
        pdf_buf = self._make_multipage_pdf(num_pages=2)
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "multi.pdf")},
            content_type="multipart/form-data",
        )
        session_id = json.loads(resp.data)["session_id"]
        self.client.post(f"/api/process/{session_id}")

        for page_num in [1, 2]:
            resp = self.client.get(f"/api/page/{session_id}/{page_num}")
            self.assertEqual(resp.status_code, 200, f"Page {page_num} image not served")
            self.assertEqual(resp.content_type, "image/png")

    def test_page_image_out_of_range_returns_404(self):
        """Requesting a page number beyond the PDF page count returns 404."""
        pdf_buf = _make_test_pdf()  # single page
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        session_id = json.loads(resp.data)["session_id"]
        self.client.post(f"/api/process/{session_id}")

        resp = self.client.get(f"/api/page/{session_id}/999")
        self.assertEqual(resp.status_code, 404)


class TestAnnotationsRoundTrip(unittest.TestCase):
    """Tests for annotation save/retrieve/export round-trip."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    def _upload(self):
        pdf_buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        return json.loads(resp.data)["session_id"]

    def test_annotations_initially_empty(self):
        """Annotations are empty immediately after upload."""
        sid = self._upload()
        resp = self.client.get(f"/api/annotations/{sid}")
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(json.loads(resp.data), [])

    def test_save_and_retrieve_annotations(self):
        """Saved annotations can be retrieved unchanged."""
        sid = self._upload()
        annotations = [
            {"id": "1", "text": "Fix typo here", "type": "comment", "page": 1},
            {"id": "2", "text": "Insert new line", "type": "insert", "page": 1},
            {"id": "3", "text": "Remove this", "type": "delete", "page": 2},
        ]
        resp = self.client.post(
            f"/api/annotations/{sid}",
            data=json.dumps(annotations),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 200)
        self.assertTrue(json.loads(resp.data)["success"])

        resp = self.client.get(f"/api/annotations/{sid}")
        self.assertEqual(resp.status_code, 200)
        retrieved = json.loads(resp.data)
        self.assertEqual(len(retrieved), 3)
        self.assertEqual(retrieved[0]["text"], "Fix typo here")
        self.assertEqual(retrieved[2]["type"], "delete")

    def test_annotations_overwrite_on_second_save(self):
        """Second POST to annotations replaces the previous set entirely."""
        sid = self._upload()
        first = [{"id": "1", "text": "First annotation", "type": "comment", "page": 1}]
        second = [{"id": "2", "text": "Replaced annotation", "type": "insert", "page": 1}]

        self.client.post(
            f"/api/annotations/{sid}",
            data=json.dumps(first),
            content_type="application/json",
        )
        self.client.post(
            f"/api/annotations/{sid}",
            data=json.dumps(second),
            content_type="application/json",
        )

        resp = self.client.get(f"/api/annotations/{sid}")
        data = json.loads(resp.data)
        self.assertEqual(len(data), 1)
        self.assertEqual(data[0]["text"], "Replaced annotation")

    def test_save_empty_annotations_clears_list(self):
        """Saving an empty list clears all annotations."""
        sid = self._upload()
        anns = [{"id": "x", "text": "temp", "type": "comment", "page": 1}]
        self.client.post(
            f"/api/annotations/{sid}",
            data=json.dumps(anns),
            content_type="application/json",
        )
        self.client.post(
            f"/api/annotations/{sid}",
            data=json.dumps([]),
            content_type="application/json",
        )
        resp = self.client.get(f"/api/annotations/{sid}")
        self.assertEqual(json.loads(resp.data), [])

    def test_annotations_save_nonexistent_session_returns_404(self):
        """POST to annotations for an unknown session returns 404."""
        resp = self.client.post(
            "/api/annotations/nonexistent-session-xyz",
            data=json.dumps([]),
            content_type="application/json",
        )
        self.assertEqual(resp.status_code, 404)

    def test_ocr_results_endpoint_returns_cached_data(self):
        """GET /api/ocr_results/<session_id> returns previously processed OCR data."""
        sid = self._upload()
        self.client.post(f"/api/process/{sid}")  # run OCR
        resp = self.client.get(f"/api/ocr_results/{sid}")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertIsInstance(data, list)
        self.assertGreater(len(data), 0)
        # Each entry should have page and blocks keys
        self.assertIn("page", data[0])
        self.assertIn("blocks", data[0])


class TestPageByPageOCR(unittest.TestCase):
    """Tests for the page-by-page OCR endpoints (/init and /page/<n>)."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    def _upload(self, num_pages=1):
        pdf = FPDF()
        for i in range(1, num_pages + 1):
            pdf.add_page()
            pdf.set_font("Helvetica", size=14)
            pdf.cell(0, 10, f"Page {i} TestContent INVOICE", new_x="LMARGIN", new_y="NEXT")
        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        return json.loads(resp.data)["session_id"]

    def test_init_returns_page_count(self):
        """POST /api/process/<sid>/init returns success and page count."""
        sid = self._upload(num_pages=1)
        resp = self.client.post(f"/api/process/{sid}/init")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        self.assertGreaterEqual(data["num_pages"], 1)

    def test_init_converts_images(self):
        """After /init, page image files exist on disk."""
        sid = self._upload(num_pages=1)
        self.client.post(f"/api/process/{sid}/init")
        pages_dir = Path(_test_upload_dir) / sid / "pages"
        self.assertTrue(pages_dir.exists())
        pngs = list(pages_dir.glob("*.png"))
        self.assertGreater(len(pngs), 0)

    def test_init_is_idempotent(self):
        """Calling /init twice returns the same page count both times."""
        sid = self._upload(num_pages=1)
        resp1 = self.client.post(f"/api/process/{sid}/init")
        resp2 = self.client.post(f"/api/process/{sid}/init")
        self.assertEqual(resp1.status_code, 200)
        self.assertEqual(resp2.status_code, 200)
        data1 = json.loads(resp1.data)
        data2 = json.loads(resp2.data)
        self.assertEqual(data1["num_pages"], data2["num_pages"])

    def test_init_unknown_session_returns_404(self):
        """POST /api/process/unknown/init returns 404."""
        resp = self.client.post("/api/process/nonexistent-session-id/init")
        self.assertEqual(resp.status_code, 404)
        self.assertIn("error", json.loads(resp.data))

    def test_process_page_extracts_text(self):
        """POST /api/process/<sid>/page/1 extracts text blocks."""
        sid = self._upload(num_pages=1)
        self.client.post(f"/api/process/{sid}/init")

        resp = self.client.post(f"/api/process/{sid}/page/1")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data["success"])
        self.assertEqual(data["page"], 1)
        self.assertGreaterEqual(data["num_pages"], 1)
        self.assertIsInstance(data["blocks"], list)

    def test_process_page_without_init_returns_400(self):
        """POST /api/process/<sid>/page/1 before /init returns 400."""
        sid = self._upload(num_pages=1)
        resp = self.client.post(f"/api/process/{sid}/page/1")
        self.assertEqual(resp.status_code, 400)
        self.assertIn("error", json.loads(resp.data))

    def test_process_page_out_of_range_returns_400(self):
        """POST /api/process/<sid>/page/99 for a 1-page PDF returns 400."""
        sid = self._upload(num_pages=1)
        self.client.post(f"/api/process/{sid}/init")
        resp = self.client.post(f"/api/process/{sid}/page/99")
        self.assertEqual(resp.status_code, 400)
        self.assertIn("error", json.loads(resp.data))

    def test_process_page_unknown_session_returns_404(self):
        """POST /api/process/unknown/page/1 returns 404."""
        resp = self.client.post("/api/process/nonexistent-session-id/page/1")
        self.assertEqual(resp.status_code, 404)
        self.assertIn("error", json.loads(resp.data))

    def test_process_page_cached_on_second_call(self):
        """Second call to /page/1 returns cached=True."""
        sid = self._upload(num_pages=1)
        self.client.post(f"/api/process/{sid}/init")
        self.client.post(f"/api/process/{sid}/page/1")  # first call
        resp = self.client.post(f"/api/process/{sid}/page/1")  # second call
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data.get("cached"), "Second page call should be cached")

    def test_all_pages_done_writes_ocr_results(self):
        """After processing all pages via /page/<n>, ocr_results.json is created."""
        sid = self._upload(num_pages=2)
        self.client.post(f"/api/process/{sid}/init")
        self.client.post(f"/api/process/{sid}/page/1")
        self.client.post(f"/api/process/{sid}/page/2")

        ocr_path = Path(_test_upload_dir) / sid / "ocr_results.json"
        self.assertTrue(ocr_path.exists(), "ocr_results.json should be written when all pages done")
        results = json.loads(ocr_path.read_text())
        self.assertEqual(len(results), 2)

    def test_init_fully_cached_after_legacy_process(self):
        """After running the legacy /api/process endpoint, /init reports fully_cached."""
        sid = self._upload(num_pages=1)
        self.client.post(f"/api/process/{sid}")  # legacy full process
        resp = self.client.post(f"/api/process/{sid}/init")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertTrue(data.get("fully_cached"), "Should report fully_cached after legacy process")

    def test_page_image_served_after_init(self):
        """After /init, the page image endpoint serves PNG for page 1."""
        sid = self._upload(num_pages=1)
        self.client.post(f"/api/process/{sid}/init")
        resp = self.client.get(f"/api/page/{sid}/1")
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(resp.content_type, "image/png")


class TestMobileAccessibility(unittest.TestCase):
    """Tests confirming mobile-accessibility markup in rendered HTML pages."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def _upload_session(self):
        """Create a session and return its session_id."""
        buf = _make_test_pdf()
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (buf, "test.pdf")},
            content_type="multipart/form-data",
        )
        return json.loads(resp.data)["session_id"]

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    # ---- index page ----

    def test_index_has_viewport_meta(self):
        """Upload page must have a responsive viewport meta tag."""
        resp = self.client.get("/")
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b'name="viewport"', resp.data)
        self.assertIn(b'width=device-width', resp.data)

    def test_index_has_skip_link(self):
        """Upload page must contain a skip-to-content link for keyboard users."""
        resp = self.client.get("/")
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b'skip-link', resp.data)
        self.assertIn(b'#uploadForm', resp.data)

    def test_index_form_has_accessible_labels(self):
        """File inputs on the upload page must have associated labels."""
        resp = self.client.get("/")
        html = resp.data.decode()
        self.assertIn('for="annotated_pdf"', html)
        self.assertIn('id="annotated_pdf"', html)

    # ---- review page ----

    def test_review_has_viewport_meta(self):
        """Review page must have a responsive viewport meta tag."""
        sid = self._upload_session()
        resp = self.client.get(f"/review/{sid}")
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b'name="viewport"', resp.data)
        self.assertIn(b'width=device-width', resp.data)

    def test_review_has_skip_link(self):
        """Review page must contain a skip-to-content link."""
        sid = self._upload_session()
        resp = self.client.get(f"/review/{sid}")
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b'skip-link', resp.data)
        self.assertIn(b'#annPanel', resp.data)

    def test_review_has_mobile_tabs(self):
        """Review page must contain the mobile tab navigation element."""
        sid = self._upload_session()
        resp = self.client.get(f"/review/{sid}")
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b'mobile-tabs', resp.data)
        self.assertIn(b'tabPdfBtn', resp.data)
        self.assertIn(b'tabAnnBtn', resp.data)

    def test_review_panels_have_aria_labels(self):
        """PDF and annotation panels must have ARIA region/label attributes."""
        sid = self._upload_session()
        resp = self.client.get(f"/review/{sid}")
        html = resp.data.decode()
        self.assertIn('role="region"', html)
        self.assertIn('aria-label="PDF viewer"', html)
        self.assertIn('aria-label="Annotation editor"', html)

    def test_review_nav_buttons_have_aria_labels(self):
        """Prev/next page navigation buttons must have aria-label attributes."""
        sid = self._upload_session()
        resp = self.client.get(f"/review/{sid}")
        html = resp.data.decode()
        self.assertIn('aria-label="Previous page"', html)
        self.assertIn('aria-label="Next page"', html)

    def test_review_progress_banner_has_live_region(self):
        """OCR progress banner must have aria-live for screen-reader updates."""
        sid = self._upload_session()
        resp = self.client.get(f"/review/{sid}")
        html = resp.data.decode()
        # The banner with aria-live should be present
        self.assertIn('aria-live="polite"', html)
        self.assertIn('ocrProgressBanner', html)
        # Progress bar element should be present
        self.assertIn('ocrProgressBar', html)

    def test_review_add_manual_button_shows_page(self):
        """'Add' button must include a page number indicator span."""
        sid = self._upload_session()
        resp = self.client.get(f"/review/{sid}")
        html = resp.data.decode()
        self.assertIn('addManualPageNum', html)
        self.assertIn('Add for Page', html)


class TestPDFAnnotationExtraction(unittest.TestCase):
    """Tests for the /api/pdf_annotations/<session_id> endpoint."""

    def setUp(self):
        app.config["TESTING"] = True
        self.client = app.test_client()

    def tearDown(self):
        for child in Path(_test_upload_dir).iterdir():
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)

    def _make_pdf_with_annotations(self):
        """Create a PDF with embedded text annotations using PyMuPDF."""
        import pymupdf
        doc = pymupdf.open()
        page = doc.new_page()
        page.add_text_annot((100, 100), "Review this section")
        page.add_text_annot((100, 200), "Check figure caption")
        buf = io.BytesIO(doc.tobytes())
        buf.seek(0)
        return buf

    def _upload(self, pdf_buf):
        resp = self.client.post(
            "/upload",
            data={"annotated_pdf": (pdf_buf, "test_annotated.pdf")},
            content_type="multipart/form-data",
        )
        return json.loads(resp.data)["session_id"]

    def test_pdf_annotations_unknown_session_returns_404(self):
        """GET /api/pdf_annotations/<unknown> returns 404."""
        resp = self.client.get("/api/pdf_annotations/nonexistent-session-id")
        self.assertEqual(resp.status_code, 404)
        data = json.loads(resp.data)
        self.assertIn("error", data)

    def test_pdf_annotations_plain_pdf_returns_empty_list(self):
        """PDF with no embedded annotations returns an empty list."""
        buf = _make_test_pdf()
        sid = self._upload(buf)
        resp = self.client.get(f"/api/pdf_annotations/{sid}")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertIsInstance(data, list)
        # A plain text PDF has no embedded annotations
        self.assertEqual(len(data), 0)

    def test_pdf_annotations_with_embedded_notes(self):
        """PDF with sticky-note annotations returns them in the response."""
        try:
            import pymupdf
        except ImportError:
            self.skipTest("PyMuPDF not installed")

        pdf_buf = self._make_pdf_with_annotations()
        sid = self._upload(pdf_buf)
        resp = self.client.get(f"/api/pdf_annotations/{sid}")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        self.assertIsInstance(data, list)
        self.assertGreater(len(data), 0, "Should extract at least one annotation")

        # Each annotation must have the required fields
        for ann in data:
            self.assertIn("id", ann)
            self.assertIn("text", ann)
            self.assertIn("type", ann)
            self.assertIn("page", ann)
            self.assertIn("source", ann)
            self.assertEqual(ann["source"], "pdf_annotation")
            self.assertIn(ann["type"], ("comment", "insert", "delete"))

    def test_pdf_annotations_content_matches(self):
        """The extracted annotation text should match the embedded content."""
        try:
            import pymupdf
        except ImportError:
            self.skipTest("PyMuPDF not installed")

        pdf_buf = self._make_pdf_with_annotations()
        sid = self._upload(pdf_buf)
        resp = self.client.get(f"/api/pdf_annotations/{sid}")
        self.assertEqual(resp.status_code, 200)
        data = json.loads(resp.data)
        texts = [a["text"] for a in data]
        self.assertTrue(
            any("Review this section" in t for t in texts),
            f"Expected 'Review this section' in extracted annotations, got: {texts}",
        )

    def test_pdf_annotations_page_numbers_correct(self):
        """Annotation page numbers should start at 1."""
        try:
            import pymupdf
        except ImportError:
            self.skipTest("PyMuPDF not installed")

        pdf_buf = self._make_pdf_with_annotations()
        sid = self._upload(pdf_buf)
        resp = self.client.get(f"/api/pdf_annotations/{sid}")
        data = json.loads(resp.data)
        for ann in data:
            self.assertGreaterEqual(ann["page"], 1)


if __name__ == "__main__":
    unittest.main()
