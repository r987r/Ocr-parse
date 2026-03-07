"""
Word document export with real Word comments and tracked changes.

Approach:
  1. Build the body content with python-docx (tracked insertions/deletions via raw XML).
  2. Inject a proper word/comments.xml part via zipfile manipulation.
  3. Update [Content_Types].xml and word/_rels/document.xml.rels accordingly.
"""

import io
import zipfile
import shutil
import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

AUTHOR = "OCR Review Parser"
INITIALS = "OCR"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_CT_COMMENTS = (
    "application/vnd.openxmlformats-officedocument"
    ".wordprocessingml.comments+xml"
)
_RT_COMMENTS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def create_word_document(annotations, original_docx_path=None, output_path="output.docx"):
    """Create a Word document with comments and tracked changes.

    Args:
        annotations : list of dicts with keys:
                        text  (str)
                        type  ('comment' | 'insert' | 'delete')
                        page  (int, optional)
        original_docx_path : path to original .docx (or None)
        output_path        : where to write the output file
    """
    date_str = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

    # ---- Step 1: build base document ----------------------------------------
    base_path = output_path + ".base.docx"
    _build_base_docx(annotations, original_docx_path, base_path, date_str)

    # ---- Step 2: inject comments.xml ----------------------------------------
    comment_anns = [a for a in annotations if a.get("type", "comment") == "comment"]
    _inject_comments(base_path, output_path, comment_anns, date_str)

    # Clean up temp file
    try:
        Path(base_path).unlink()
    except OSError:
        pass

    return output_path


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _build_base_docx(annotations, original_docx_path, out_path, date_str):
    """Build the main document body using python-docx."""

    if original_docx_path and Path(original_docx_path).exists():
        doc = Document(original_docx_path)
        # Append a separator
        sep = doc.add_paragraph()
        sep.add_run("─" * 50)
        doc.add_paragraph()
        intro = doc.add_paragraph()
        intro.add_run("Annotations from PDF Review").bold = True
        doc.add_paragraph()
    else:
        doc = Document()
        doc.add_heading("Reviewed Document", level=1)
        doc.add_paragraph(
            "This document was generated from PDF annotations by OCR Review Parser."
        )
        doc.add_paragraph()

    comment_id_counter = [0]

    for ann in annotations:
        ann_type = ann.get("type", "comment")
        text = ann.get("text", "").strip()
        page = ann.get("page")
        label = f"(Page {page}) " if page else ""

        if not text:
            continue

        if ann_type == "comment":
            # Add a placeholder paragraph; we embed a unique marker text
            # that _inject_comments will replace with proper comment XML.
            para = doc.add_paragraph()
            marker = f"__OCR_COMMENT_{comment_id_counter[0]}__"
            run = para.add_run(marker)
            run.italic = True
            comment_id_counter[0] += 1

        elif ann_type == "insert":
            # Tracked insertion
            para = doc.add_paragraph()
            _add_tracked_insertion(para, f"{label}{text}", date_str, _next_rev_id())

        elif ann_type == "delete":
            # Tracked deletion
            para = doc.add_paragraph()
            _add_tracked_deletion(para, f"{label}{text}", date_str, _next_rev_id())

    doc.save(out_path)


_rev_counter = [200]


def _next_rev_id():
    _rev_counter[0] += 1
    return _rev_counter[0]


def _add_tracked_insertion(para, text, date_str, rev_id):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), AUTHOR)
    ins.set(qn("w:date"), date_str)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    ins.append(r)
    para._p.append(ins)


def _add_tracked_deletion(para, text, date_str, rev_id):
    d = OxmlElement("w:del")
    d.set(qn("w:id"), str(rev_id))
    d.set(qn("w:author"), AUTHOR)
    d.set(qn("w:date"), date_str)
    r = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.text = text
    dt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(dt)
    d.append(r)
    para._p.append(d)


# ---------------------------------------------------------------------------
# ZIP-level comment injection
# ---------------------------------------------------------------------------

def _inject_comments(base_path, output_path, comment_anns, date_str):
    """Open the base docx, inject comments.xml, update metadata, write output."""

    with open(base_path, "rb") as f:
        in_bytes = f.read()

    in_buf = io.BytesIO(in_bytes)
    out_buf = io.BytesIO()

    # Build comment list with IDs
    comments = [
        {
            "id": idx,
            "text": ann.get("text", "").strip(),
            "page": ann.get("page"),
        }
        for idx, ann in enumerate(comment_anns)
        if ann.get("text", "").strip()
    ]

    has_comments = bool(comments)

    with zipfile.ZipFile(in_buf, "r") as zin, zipfile.ZipFile(
        out_buf, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "word/document.xml":
                data = _process_document_xml(data, comments, date_str)

            elif item.filename == "[Content_Types].xml" and has_comments:
                data = _add_ct_comments(data)

            elif item.filename == "word/_rels/document.xml.rels" and has_comments:
                data = _add_rel_comments(data)

            zout.writestr(item, data)

        if has_comments:
            zout.writestr("word/comments.xml", _build_comments_xml(comments, date_str))

    with open(output_path, "wb") as f:
        f.write(out_buf.getvalue())


def _process_document_xml(data, comments, date_str):
    """Find annotation placeholders and replace them with proper comment markup."""
    if not comments:
        return data

    try:
        tree = etree.fromstring(data)
    except etree.XMLSyntaxError:
        return data

    W = _W

    # Build map: marker string -> comment dict
    comment_map = {f"__OCR_COMMENT_{c['id']}__": c for c in comments}

    # Find all text elements (<w:t>) containing a marker
    for t_el in tree.iter(f"{{{W}}}t"):
        text_val = t_el.text or ""
        if text_val not in comment_map:
            continue

        c = comment_map[text_val]
        cid = c["id"]

        # The structure is: <w:p> ... <w:r><w:t>MARKER</w:t></w:r> ...
        r_el = t_el.getparent()
        if r_el is None or r_el.tag != f"{{{W}}}r":
            continue

        # Replace the marker text with a readable label
        label = f"[Annotation \u2013 Page {c['page']}]" if c.get("page") else "[Annotation]"
        t_el.text = label

        parent = r_el.getparent()
        if parent is None:
            continue

        # Insert commentRangeStart before the run (recalculate index each time)
        crs = etree.Element(f"{{{W}}}commentRangeStart")
        crs.set(f"{{{W}}}id", str(cid))
        idx = list(parent).index(r_el)
        parent.insert(idx, crs)

        # Insert commentRangeEnd after the run (recalculate after previous insert)
        run_idx = list(parent).index(r_el)
        cre = etree.Element(f"{{{W}}}commentRangeEnd")
        cre.set(f"{{{W}}}id", str(cid))
        parent.insert(run_idx + 1, cre)

        # Insert commentReference run after commentRangeEnd (recalculate again)
        cre_idx = list(parent).index(cre)
        ref_run = etree.Element(f"{{{W}}}r")
        rpr = etree.SubElement(ref_run, f"{{{W}}}rPr")
        rst = etree.SubElement(rpr, f"{{{W}}}rStyle")
        rst.set(f"{{{W}}}val", "CommentReference")
        cr = etree.SubElement(ref_run, f"{{{W}}}commentReference")
        cr.set(f"{{{W}}}id", str(cid))
        parent.insert(cre_idx + 1, ref_run)

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def _build_comments_xml(comments, date_str):
    W = _W
    # Use nsmap to declare the 'w' prefix properly
    root = etree.Element(f"{{{W}}}comments", nsmap={"w": W})
    for c in comments:
        comment_el = etree.SubElement(root, f"{{{W}}}comment")
        comment_el.set(f"{{{W}}}id", str(c["id"]))
        comment_el.set(f"{{{W}}}author", AUTHOR)
        comment_el.set(f"{{{W}}}date", date_str)
        comment_el.set(f"{{{W}}}initials", INITIALS)

        para_el = etree.SubElement(comment_el, f"{{{W}}}p")

        # Optional page reference run
        if c.get("page"):
            # Add page label as first run
            page_run = etree.SubElement(para_el, f"{{{W}}}r")
            page_rpr = etree.SubElement(page_run, f"{{{W}}}rPr")
            etree.SubElement(page_rpr, f"{{{W}}}b")
            page_t = etree.SubElement(page_run, f"{{{W}}}t")
            page_t.text = f"[Page {c['page']}] "
            page_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        run_el = etree.SubElement(para_el, f"{{{W}}}r")
        t_el = etree.SubElement(run_el, f"{{{W}}}t")
        t_el.text = c["text"]
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_ct_comments(data):
    """Add the comments Override entry to [Content_Types].xml."""
    NS = "http://schemas.openxmlformats.org/package/2006/content-types"
    try:
        tree = etree.fromstring(data)
    except etree.XMLSyntaxError:
        return data

    existing = tree.findall(f"{{{NS}}}Override")
    if not any(
        o.get("PartName", "") == "/word/comments.xml" for o in existing
    ):
        override = etree.SubElement(tree, f"{{{NS}}}Override")
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", _CT_COMMENTS)

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_rel_comments(data):
    """Add the comments Relationship to word/_rels/document.xml.rels."""
    NS = "http://schemas.openxmlformats.org/package/2006/relationships"
    try:
        tree = etree.fromstring(data)
    except etree.XMLSyntaxError:
        return data

    existing = tree.findall(f"{{{NS}}}Relationship")
    if not any(r.get("Type", "").endswith("/comments") for r in existing):
        rel = etree.SubElement(tree, f"{{{NS}}}Relationship")
        # Pick an unused rId
        used_ids = {r.get("Id", "") for r in existing}
        rid = "rIdOcrComments"
        counter = 1
        while rid in used_ids:
            rid = f"rIdOcrComments{counter}"
            counter += 1
        rel.set("Id", rid)
        rel.set("Type", _RT_COMMENTS)
        rel.set("Target", "comments.xml")

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
